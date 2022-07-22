# Главный файл проекта

# подключаем библиотеки

import os
import docx
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Font, Border, Side, Alignment, PatternFill
from openpyxl.styles.numbers import BUILTIN_FORMATS

# подключаем модули

from doc_to_docx import doc_to_docx

# устанавливаем переменные для работы
PATH = './files/ЖЕТЫБАЙ'                                              # путь к папке или файлу с данными
BN_PATH = './files/Перечень зданий. 1 этап. Ред 1.xlsx'               # путь к файлу с наименованием зданий
SAVE_PATH = './files/'                                                # путь к папке для сохранения результата

# формируем таблицу с названиями зданий

building_names = pd.read_excel(BN_PATH, sheet_name='Список зданий', header=0, engine='openpyxl')

# формируем перечень файлов для извлечения данных
paths = []
folder = PATH

# формируем список файлов

for root, dirs, files in os.walk(folder):
    for file in files:
        if file.endswith('docx') and not file.startswith('~') and (os.path.splitext(file)[0]=='Вед. смонт. обор общ.'):
            paths.append(os.path.join(root, file))

# обработка файлов ведомости смонтированного оборудования
# подключение к файлу, считывание информации

for file in paths:
    doc = docx.Document(file)
    ob_name = doc.paragraphs[4].text[8:]  # запомним имя объекта

    # проверяем исключение (второй вариант написания наименования объекта)

    if ob_name == 'ванного оборудования АПС' or ob_name == 'ванного оборудования АПС на':
        ob_name = doc.paragraphs[5].text[8:]  # запомним имя объекта

    # приводим имя объекта к стандарту
    ob_name = ob_name.replace('"','')
    ob_name = ob_name.replace('/', '')
    ob_name = ob_name.replace(' -', '-')
    ob_name = ob_name.replace(' АО ', '')
    ob_name = ob_name.replace('ММГ', '')
    ob_name = ob_name.replace(chr(187), '')
    ob_name = ob_name.replace(chr(171), '')
    ob_name = ob_name.replace('ЦПТГ и ЭГХ', 'ЦПТГиЭГХ')
    ob_name = ob_name.replace('ЦДНГ-2 ЖМГ', 'ЦДНГ-2 ПУ ЖМГ')
    ob_name = ob_name.lstrip()
    ob_name = ob_name[:70]


    # проведем поиск в таблице с наименованием зданий
    if len(building_names.loc[building_names['Наименование объектов полное'] == ob_name, '№ п/п']) > 0:
        ob_number = building_names.loc[building_names['Наименование объектов полное'] == ob_name, '№ п/п'].values[0]
    else:
        ob_number = False

    if ob_number:
        ob_name = str(ob_number).rjust(3, '0')+'_'+ob_name
    else:
        ob_name = 'xxx_' + ob_name

    # обработка таблиц

    cell_text = []

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text.append(cell.text)

    # формируем список для итоговой таблицы файла

    cell_list = []

    for i in range(0,len(cell_text),6):
        if (cell_text[i+1]!='Наименование оборудования') and \
                (cell_text[i + 1] != '\nНаименование оборудования') and \
                (cell_text[i + 1] != 'Наименование оборудования\n') and \
                (cell_text[i+1]!='\nНаименование оборудования\n') and \
                (cell_text[i + 1] != 'Наименование \nоборудования') and \
                (cell_text[i + 1] != '\nНаименование\nоборудования\n') and \
                (cell_text[i+1]!='2'):

            # if (cell_text[i + 1].replace('\n','') != 'Наименование оборудования') and (cell_text[i + 1] != '2'):
                cell_list.append([cell_text[i+1], cell_text[i+2], cell_text[i+3], cell_text[i+4]])

    # формируем итоговую таблицу для файла

    data = pd.DataFrame(cell_list, columns=['name', 'short_name', 'unit', 'amount'], dtype='string')
    print(ob_name)
    data['amount'] = data['amount'].astype('int')

    # записываем итоговую таблицу в файл

    data.to_excel(SAVE_PATH+ob_name+'.xlsx', header=['Наименование', 'Тип', 'Ед. изм', 'Кол-во'], \
                  sheet_name='Спецификация', index = False, engine='openpyxl')

    # преобразуем файл Excel
    # открываем файл для изменения стиля ячеек

    wb = load_workbook(SAVE_PATH+ob_name+'.xlsx')
    sheet = wb['Спецификация']

    # фиксируем стиль для границ ячейки (рамку)

    bd = Side(style='thin', color="000000")

    # проводим перебор ячеек и меняем стили

    col_name = 'ABCD'

    for k in range(0, 4):
        for i in range(1, len(data) + 2):
            sheet[col_name[k] + str(i)].font = Font(name='Calibri', bold=False, size=11)
            sheet[col_name[k] + str(i)].border = Border(left=bd, top=bd, right=bd, bottom=bd)
            # sheet[col_name[k] + str(i)].alignment = Alignment(wrap_text=True)

            # изменение формата ячейки для строковых значений

            if k == 0:
                sheet[col_name[k] + str(i)].alignment = Alignment(horizontal='left', vertical='center', \
                                                                  wrap_text=True)

            # изменение формата ячейки для строковых значений

            if (k == 1) or (k == 2):
                sheet[col_name[k] + str(i)].alignment = Alignment(horizontal='center', vertical='center', \
                                                                  wrap_text=True)
            # изменение формата ячейки для целых чисел

            if k == 3:
                sheet[col_name[k] + str(i)].number_format = BUILTIN_FORMATS[1]
                sheet[col_name[k] + str(i)].alignment = Alignment(horizontal='center', vertical='center', \
                                                                  wrap_text=True)

            # изменение формата ячейки для денежных значений

            # if (k == 5) or (k == 6):
            #     sheet[col_name[k] + str(i)].number_format = BUILTIN_FORMATS[4]
            #     sheet[col_name[k] + str(i)].alignment = Alignment(horizontal='right', vertical='center', wrap_text=True)

            # изменение формата ячеек первой строки

            if (i == 1):
                sheet[col_name[k] + str(i)].fill = PatternFill("solid", fgColor="00AAEE")
                sheet[col_name[k] + str(i)].alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                sheet[col_name[k] + str(i)].font = Font(bold=True)

    # изменяем ширину столбцов

    sheet.column_dimensions['A'].width = 40
    sheet.column_dimensions['B'].width = 25
    sheet.column_dimensions['C'].width = 10
    sheet.column_dimensions['D'].width = 10

    # записываем результат в файл

    wb.save(SAVE_PATH+ob_name+'.xlsx')
